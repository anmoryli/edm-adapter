from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "report" / "ei_paper" / "edm_struct_lora_ei_paper.docx"
ASSETS = ROOT / "report" / "paper_assets"


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_font(paragraph, size: float = 10.5) -> None:
    for run in paragraph.runs:
        set_run_font(run, size=size)


def paragraph(doc: Document, text: str = "", style: str | None = None, align=None, size: float = 10.5, bold: bool = False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0.74) if style is None else None
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=14 if level == 1 else 12, bold=True)
    return p


def equation(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=9.5, bold=bold)


def set_cell_width(cell, width_cm: float):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.first_child_found_in("w:tcW")
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")


def set_table_borders(table, header_rows: int = 1) -> None:
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                element = borders.find(qn(f"w:{edge}"))
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    borders.append(element)
                element.set(qn("w:val"), "nil")
            if row_idx == 0:
                top = borders.find(qn("w:top"))
                top.set(qn("w:val"), "single")
                top.set(qn("w:sz"), "12")
                top.set(qn("w:color"), "000000")
            if row_idx == header_rows - 1:
                bottom = borders.find(qn("w:bottom"))
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "8")
                bottom.set(qn("w:color"), "000000")
            if row_idx == len(table.rows) - 1:
                bottom = borders.find(qn("w:bottom"))
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "12")
                bottom.set(qn("w:color"), "000000")


def table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths: list[float]):
    cap = paragraph(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5, bold=True)
    cap.paragraph_format.first_line_indent = None
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for idx, h in enumerate(headers):
        set_cell_text(t.rows[0].cells[idx], h, bold=True)
        set_cell_width(t.rows[0].cells[idx], widths[idx])
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(value) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[idx], value, align=align)
            set_cell_width(cells[idx], widths[idx])
    set_table_borders(t)
    paragraph(doc, "", size=4)
    return t


def figure(doc: Document, filename: str, caption_text: str):
    path = ASSETS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(5.7))
    cap = paragraph(doc, caption_text, align=WD_ALIGN_PARAGRAPH.CENTER, size=9.5, bold=True)
    cap.paragraph_format.first_line_indent = None


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("• " + text)
    set_run_font(r, size=10.5)
    return p


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10.5)

    title = paragraph(doc, "EDM-StructLoRA：面向电子音乐生成的结构与属性感知参数高效适配方法", align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
    title.paragraph_format.first_line_indent = None
    author = paragraph(doc, "作者姓名  单位名称  邮箱", align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
    author.paragraph_format.first_line_indent = None

    heading(doc, "摘要", 1)
    paragraph(doc, "文本生成音乐模型已经能够根据自然语言提示生成完整音频片段，但在电子舞曲（Electronic Dance Music, EDM）场景中仍面临结构控制不足的问题。EDM 的创作依赖明确的段落组织、稳定的 BPM、可感知的能量变化、低频冲击和循环边界，而通用文本提示往往难以稳定约束这些时间相关属性。针对这一问题，本文提出 EDM-StructLoRA，一种面向 EDM 生成的结构与属性感知参数高效适配框架。该方法基于 ACE-Step 音乐生成基础模型，将传统单一 LoRA 扩展为共享 EDM LoRA、段落专家 LoRA、能量专家 LoRA 和子风格专家 LoRA，并通过元数据路由器在训练与推理阶段动态计算 adapter mixture。进一步地，本文构建 latent-frame aligned control curve，将 section、subgenre、energy、BPM、beat phase、low-frequency ratio、onset density、loop boundary 和标签置信度编码为 36 维帧级控制信号，再通过轻量 MLP control conditioner 转换为额外文本条件 token 注入扩散 Transformer。实验原型构建了 231 个原始音频源、203 个有效处理源和 6,016 条 8 秒 EDM 训练片段，生成对应的 ACE-Step latents、text tokens 和控制曲线，并完成 HuggingFace Dataset 转换、LoRA bundle 保存和可控推理流程。与普通 LoRA 微调相比，本文方法的核心优势在于将 EDM 领域结构知识显式纳入参数高效适配过程，为 BPM、段落、能量与子风格可控的文本生成音乐提供了一条可复现的工程路径。")
    paragraph(doc, "关键词：文本生成音乐；电子舞曲；LoRA；可控生成；Adapter Routing；Control Token；ACE-Step", size=10.5, bold=True)

    heading(doc, "Abstract", 1)
    paragraph(doc, "Text-to-music models can synthesize musical audio from natural-language prompts, but they still lack reliable structural control for electronic dance music (EDM). EDM production depends on explicit sections, stable BPM, energy progression, bass impact, onset density and loop boundaries, which are difficult to constrain by text prompts alone. This paper presents EDM-StructLoRA, a structure- and attribute-aware parameter-efficient adaptation framework for controllable EDM generation. Built on the ACE-Step foundation model, EDM-StructLoRA replaces a single fixed LoRA adapter with a shared EDM adapter, section experts, energy experts and subgenre experts. A metadata-conditioned router dynamically activates adapter mixtures during training and inference. In addition, a latent-frame aligned control curve encodes section, subgenre, energy, BPM, beat phase, low-frequency ratio, onset density, loop markers and confidence values into 36-dimensional frame-level controls, which are projected into additional text-conditioning tokens. The prototype builds 6,016 cleaned 8-second EDM clips from 231 raw sources, caches ACE-Step latents, constructs control assets and implements a full training-inference pipeline.")
    paragraph(doc, "Keywords: text-to-music generation; EDM; LoRA; controllable generation; adapter routing; control token; ACE-Step", size=10.5, bold=True)

    heading(doc, "一、引言", 1)
    paragraph(doc, "近年来，文本生成音频与文本生成音乐模型快速发展。MusicLM、MusicGen、AudioLDM 以及 ACE-Step 等模型表明，基于大规模音频数据和强表征模型的生成系统已经能够根据自然语言描述合成具有较高可听性的音乐片段。然而，音乐生成任务并不只是根据风格词生成声音。对于电子舞曲，尤其是 festival EDM、melodic house、progressive house、future bass 等类型，生成质量很大程度上取决于段落功能、节奏速度、能量轨迹、低频分布和循环可用性。")
    paragraph(doc, "现有通用文本生成音乐模型通常将用户意图压缩进 prompt。该方式在描述风格、乐器和情绪时较为有效，但对 EDM 的结构化属性控制仍不稳定。首先，模型可能无法稳定区分 intro、build-up、drop、breakdown 和 loop 等段落。其次，BPM、鼓点密度、低频能量和 loop 边界是时间相关信号，单纯写入 caption 并不能保证在生成过程中被精确执行。第三，普通 LoRA 微调一般只训练一个固定 adapter，它可以学习某个领域的整体音色偏移，却难以对不同段落、能量等级和子风格进行差异化建模。")
    paragraph(doc, "为解决上述问题，本文提出 EDM-StructLoRA。其核心思想是：将 EDM 生成中的领域适配从单一风格迁移扩展为结构与属性条件下的动态参数适配。模型使用一个共享 LoRA 捕获 EDM 的整体域偏移，同时使用多个低秩专家 LoRA 表征不同 section、energy 和 subgenre。router 根据样本元数据计算 adapter 权重，使模型在处理不同输入时激活不同参数子空间。同时，本文设计 36 维 latent 对齐控制曲线，将 BPM、beat phase、onset density、low-frequency ratio 和 loop markers 等时间信息显式注入文本条件序列。")
    bullet(doc, "提出面向 EDM 生成的结构与属性感知 LoRA 路由框架，将固定 LoRA 扩展为共享 adapter 与多专家 adapter 的动态组合。")
    bullet(doc, "设计 latent-frame aligned EDM control curve，将 section、subgenre、energy、BPM、节拍相位、低频比例、onset 密度、loop 边界与置信度编码为统一控制张量。")
    bullet(doc, "实现控制曲线到文本条件 token 的轻量投影模块，使结构控制以额外 conditioning tokens 的形式注入 ACE-Step 扩散 Transformer。")
    bullet(doc, "构建包含 6,016 条 8 秒训练片段的 EDM 数据处理与训练流水线，覆盖音频切片、标签构建、caption 生成、latent 缓存、HuggingFace Dataset 转换、adapter bundle 保存与可控推理。")

    heading(doc, "二、相关工作", 1)
    heading(doc, "2.1 文本生成音乐", 2)
    paragraph(doc, "文本生成音乐模型的研究目标是根据自然语言描述生成符合语义和音乐风格的音频。MusicLM 将文本到音乐生成表述为层次化序列到序列任务，MusicGen 使用压缩离散音频 token 并通过单个语言模型建模多流表示，AudioLDM 将潜空间扩散模型引入文本到音频生成，ACE-Step 则进一步面向音乐生成基础模型构建开放式训练与推理框架。尽管这些模型提升了生成音质和文本相关性，但它们通常不直接面向 EDM 的段落结构进行建模。")
    heading(doc, "2.2 可控音乐生成", 2)
    paragraph(doc, "可控音乐生成关注如何将节奏、和弦、旋律、音色、参考音频或时间变化属性作为条件注入生成模型。MusiConGen 等工作表明，单纯文本条件不足以精确控制 rhythm 和 chord 等音乐属性，引入自动提取或用户指定的时间条件能够提高生成一致性。ControlNet 在图像扩散模型中展示了保留基础模型能力并引入额外控制分支的思路，该思想对音频生成中的结构控制具有启发意义。")
    heading(doc, "2.3 参数高效微调与 LoRA", 2)
    paragraph(doc, "LoRA 通过冻结预训练模型参数并在权重矩阵中加入低秩增量，显著降低大模型微调的训练参数量。给定预训练权重 W₀ ∈ R^(d×k)，LoRA 将下游适配表示为：")
    equation(doc, "W = W₀ + ΔW,    ΔW = (α/r)BA")
    paragraph(doc, "普通 LoRA 通常对应一个固定的 ΔW，适合学习整体域偏移，但在结构化音乐生成中难以表达同一基础模型在不同段落和属性下应使用不同适配方向的需求。本文将 LoRA 进一步组织成共享 adapter 和专家 adapter 的集合，并通过 router 动态计算组合权重。")

    heading(doc, "三、问题定义与需求分析", 1)
    paragraph(doc, "给定文本提示 p、EDM 控制元数据 m 和基础音乐生成模型 Gθ，本文目标是生成音频片段 x，使其同时满足文本语义、EDM 子风格、段落功能、目标 BPM、能量等级和 loop 结构约束。元数据可表示为：")
    equation(doc, "m = {s, e, g, b, q, c}")
    paragraph(doc, "其中 s 表示 section，e 表示 energy，g 表示 subgenre，b 表示 BPM，q 表示样本质量权重，c 表示标签置信度。从应用角度看，该问题的核心痛点包括段落功能不稳定、BPM 和节奏控制弱、能量标签表达粗糙、单一 LoRA 适配能力有限以及数据工程复杂。")

    heading(doc, "四、方法", 1)
    heading(doc, "4.1 总体框架", 2)
    paragraph(doc, "EDM-StructLoRA 的输入由文本提示和 EDM 控制元数据组成。router 根据 section、energy、subgenre、BPM 和置信度计算 LoRA 专家权重，动态激活共享 LoRA 与专家 LoRA。与此同时，控制曲线构建器生成与 latent frame 对齐的 36 维控制序列，control conditioner 将该序列压缩为固定数量的控制 token，并拼接到文本编码器输出之后。")
    figure(doc, "fig1_architecture.png", "图 1  EDM-StructLoRA 总体结构示意图")
    heading(doc, "4.2 结构与属性感知 LoRA 路由", 2)
    paragraph(doc, "本文将 LoRA adapter 集合定义为共享 EDM adapter、section 专家集合、energy 专家集合和 subgenre 专家集合。当前实现中包含 1 个共享 adapter、9 个 section experts、4 个 energy experts 和 11 个 subgenre experts，共 25 个 adapter。")
    equation(doc, "A = {a₀} ∪ A_sec ∪ A_ene ∪ A_sub")
    paragraph(doc, "对于 batch 中的样本 i，router 根据元数据计算权重，batch 级 adapter 权重为样本权重的均值。工程实现中，共享 adapter 权重固定为 1.0，section、energy 和 subgenre 专家的权重分别由标签置信度、样本权重和 BPM 偏移进行缩放。")
    equation(doc, "ω_sec = λ_s c_s q (1 + λ_b |b - 128| / 60)")
    equation(doc, "ω_ene = λ_e c_e q (1 + λ_b |b - 128| / 60)")
    equation(doc, "ω_sub = λ_g c_g q")
    paragraph(doc, "该路由机制的作用不是替代文本 prompt，而是在参数空间中提供与 EDM 结构相关的适配方向。例如，当输入指定为 high energy melodic house drop 时，系统会同时激活 shared EDM adapter、section_drop、energy_high 和 subgenre_melodic_house。")
    heading(doc, "4.3 Latent 对齐控制曲线", 2)
    paragraph(doc, "为了避免模型只依赖 caption 理解结构属性，本文构建 latent-frame aligned control curve。对于每个音频片段，系统生成控制矩阵 C ∈ R^(T×F)，其中 T 为 latent frame 数，F=36 为控制特征维度。控制特征包括 section one-hot、subgenre one-hot、energy one-hot、energy scalar、normalized BPM、BPM confidence、beat phase sin/cos、time position、low-frequency ratio、onset density、loop start/end markers、tag confidence mean 和 quality weight。")
    equation(doc, "C = [c₁, c₂, ..., c_T] ∈ R^(T×36)")
    equation(doc, "φ_t = 2π(b/60)τ_t,    c_t^phase = [sin(φ_t), cos(φ_t)]")
    figure(doc, "fig3_control_curve.png", "图 2  latent 对齐 EDM 控制曲线示例")
    heading(doc, "4.4 控制 token 投影器", 2)
    paragraph(doc, "控制曲线 C 不能直接输入文本编码序列，因此本文设计轻量 MLP control conditioner。首先对时间维进行自适应池化，将 T 帧控制曲线压缩为 K 个控制位置，随后使用 LayerNorm、Linear、SiLU、Dropout 和 Linear 投影到文本 embedding 维度。当前实现中 F=36、K=8、d=768、hidden dimension 为 512。")
    equation(doc, "C̃ = Pool(C) ∈ R^(K×F)")
    equation(doc, "Z_c = MLP(C̃) + E_type,    H' = [H_text; Z_c]")
    heading(doc, "4.5 训练目标", 2)
    paragraph(doc, "本文基于 ACE-Step 的扩散训练流程进行适配。基础模型主体参数保持冻结，仅训练 LoRA 参数和 control conditioner 参数，从而实现参数高效适配。训练目标可写为加权均方误差：")
    equation(doc, "L = E[ w_i || ε - ε_{θ,Δθ(m_i)}(z_t, t, H'_i) ||² ]")

    heading(doc, "五、系统实现", 1)
    paragraph(doc, "项目实现了从原始音频到训练样本的完整流水线。处理过程包括：原始音频收集、音频切片、EDM 元数据构建、音频特征分析、caption 生成、数据集划分、ACE-Step latent 缓存、text token 缓存、control curve 生成和 HuggingFace Dataset 转换。每条样本包含音频路径、caption、section、energy、subgenre、BPM、质量分数、标签置信度、latent path、control path 和 text token path。")
    figure(doc, "fig4_pipeline.png", "图 3  EDM 数据处理与训练流水线")
    table(
        doc,
        "表 1  EDM 数据集规模统计",
        ["项目", "数值"],
        [
            ["原始音频源文件数", "231"],
            ["有效处理源文件数", "203"],
            ["训练质量片段数", "6,016"],
            ["训练集 / 验证集 / 测试集", "4,833 / 557 / 626"],
            ["单片段时长", "8 s"],
            ["ACE-Step latents", "6,016"],
            ["控制曲线", "6,016"],
            ["控制特征维度", "36"],
        ],
        [7.2, 5.0],
    )
    figure(doc, "fig2_dataset_distribution.png", "图 4  EDM 数据集标签分布")
    table(
        doc,
        "表 2  EDM-StructLoRA 关键配置",
        ["模块", "配置"],
        [
            ["基础模型", "ACE-Step-v1-3.5B"],
            ["共享 LoRA", "rank=32, alpha=16, rsLoRA"],
            ["专家 LoRA", "rank=8, alpha=8, rsLoRA"],
            ["Adapter 数量", "25"],
            ["Router scale", "section=0.90, energy=0.35, subgenre=0.45, BPM=0.20"],
            ["Control conditioner", "feature dim=36, token count=8, hidden dim=512"],
            ["训练精度", "bf16-mixed"],
            ["推荐学习率", "1×10⁻⁴"],
            ["梯度累积", "8"],
            ["采样策略", "confidence-aware weighted sampling"],
        ],
        [5.0, 9.5],
    )
    paragraph(doc, "训练过程保存的不再是单个 LoRA 文件，而是包含 manifest、control conditioner 和多个 adapter 权重的 bundle。推理时，系统加载基础 ACE-Step 模型和 LoRA bundle，根据用户输入的 section、energy、subgenre 和 BPM 计算路由权重，并注入控制 token 进行生成。生成完成后，系统同时保存 wav 音频和 JSON sidecar，用于记录 prompt、adapter 权重、控制参数和输出路径。")

    heading(doc, "六、实验设计与原型验证", 1)
    table(
        doc,
        "表 3  消融实验设计",
        ["编号", "系统", "实验目的"],
        [
            ["A0", "ACE-Step base", "未适配基线"],
            ["A1", "Plain LoRA r=64", "参数高效领域适配基线"],
            ["A2", "Metadata-rich caption LoRA", "验证 caption 工程是否足够"],
            ["A3", "Section-aware LoRA", "验证段落专家 adapter 的作用"],
            ["A4", "Section + attribute routed LoRA", "验证动态路由机制"],
            ["A5", "Full EDM-StructLoRA", "验证 router + control token + weighted sampling"],
        ],
        [1.5, 5.2, 7.5],
    )
    table(
        doc,
        "表 4  建议评价指标",
        ["指标", "含义"],
        [
            ["BPM Error", "目标 BPM 与生成音频估计 BPM 的绝对误差"],
            ["Onset Density Error", "目标鼓点密度与生成鼓点密度之间的差异"],
            ["Low-frequency Ratio", "低频能量占比，用于衡量 drop 或 bass-heavy 片段表现"],
            ["Loop Similarity", "loop 首尾频谱或 embedding 相似度，用于衡量循环可用性"],
            ["Section Controllability", "目标 section 与分类器预测 section 的一致性"],
            ["Text-audio Similarity", "文本提示与生成音频 embedding 的相似度"],
            ["FAD / Embedding Distance", "生成分布与参考 EDM 数据分布之间的距离"],
            ["Human Blind Test", "风格匹配、drop 冲击、loop 可用性和总体音质评分"],
        ],
        [5.0, 9.0],
    )
    paragraph(doc, "当前原型已完成方法实现和系统级验证：完成 6,016 条 EDM 片段的 metadata、caption、latent、control curve 和 text token 构建，control assets 覆盖率为 6,016/6,016；完成 train/validation/test 数据划分，并将训练集转换为 HuggingFace Dataset 格式，训练样本数为 4,833；实现动态 LoRA router，可根据 batch metadata 输出 shared、section、energy 和 subgenre adapter 权重；实现 control conditioner，将 36 维帧级控制序列压缩为 8 个文本条件 token，并拼接至 UMT5 embedding；完成 step=1000 的 EDM control LoRA bundle 保存，bundle 内包含 25 个 adapter 权重、manifest 和 control conditioner；实现可控推理脚本，可根据用户输入的 prompt、section、energy、subgenre 和 BPM 动态加载 adapter mixture 并生成音频。")
    paragraph(doc, "需要说明的是，本文当前阶段重点是方法设计、系统实现和可复现实验框架构建。若用于正式投稿，还需要补充完整主观听测、自动指标和多基线消融结果。本文已经在实验设计层面给出了可直接执行的比较方案和指标体系，后续可在同一代码框架下扩展最终量化结果。")

    heading(doc, "七、讨论", 1)
    paragraph(doc, "EDM-StructLoRA 的主要优势在于将 EDM 音乐制作知识显式融入模型适配过程。普通 LoRA 学习的是平均域偏移，容易把不同段落、能量和子风格混合为单一风格特征。本文通过专家 adapter 和 router 将结构差异映射到不同低秩参数子空间，使模型能够在生成时根据条件进行参数级切换。另一方面，control curve 将 BPM、beat phase、onset density 和 loop boundary 等时间信号显式注入条件序列，弥补了 prompt 对连续控制表达不足的问题。")
    paragraph(doc, "当前原型仍存在三方面局限。第一，数据分布存在不均衡，drop 和 high energy 样本占比较高，breakdown 等少数段落样本不足。第二，本文采用启发式 router scale，虽然便于解释和调试，但尚未学习到最优路由函数。第三，当前控制 token 使用平均池化压缩控制曲线，可能损失部分细粒度时间变化信息。未来可探索可学习 temporal attention、router 正则化和类别均衡采样，以进一步提高控制精度。")

    heading(doc, "八、结论", 1)
    paragraph(doc, "本文提出 EDM-StructLoRA，一种面向 EDM 文本生成音乐的结构与属性感知参数高效适配方法。该方法基于 ACE-Step 音乐生成基础模型，将 LoRA 从单一 adapter 扩展为共享 EDM adapter、section experts、energy experts 和 subgenre experts，并通过元数据路由器动态激活 adapter mixture。同时，本文构建 latent 对齐控制曲线，将 EDM 的段落、BPM、节拍相位、低频能量、onset 密度、loop 边界和置信度编码为显式控制信号，再投影为控制 token 注入文本条件序列。项目实现了包含 6,016 条 EDM 片段的数据处理、训练、保存和推理闭环，为可控 EDM 生成提供了可复现的研究原型。未来工作将围绕完整量化评测、人工盲听实验、可学习 router 和更细粒度时间控制展开。")

    heading(doc, "参考文献", 1)
    refs = [
        "E. J. Hu, Y. Shen, P. Wallis, Z. Allen-Zhu, Y. Li, S. Wang, L. Wang, and W. Chen, “LoRA: Low-Rank Adaptation of Large Language Models,” arXiv:2106.09685, 2021.",
        "J. Copet, F. Kreuk, I. Gat, T. Remez, D. Kant, G. Synnaeve, Y. Adi, and A. Défossez, “Simple and Controllable Music Generation,” arXiv:2306.05284, 2023.",
        "A. Agostinelli et al., “MusicLM: Generating Music From Text,” arXiv:2301.11325, 2023.",
        "H. Liu et al., “AudioLDM: Text-to-Audio Generation with Latent Diffusion Models,” arXiv:2301.12503, 2023.",
        "L. Zhang, A. Rao, and M. Agrawala, “Adding Conditional Control to Text-to-Image Diffusion Models,” ICCV, 2023.",
        "Y.-H. Lan, W.-Y. Hsiao, H.-C. Cheng, and Y.-H. Yang, “MusiConGen: Rhythm and Chord Control for Transformer-Based Text-to-Music Generation,” arXiv:2407.15060, 2024.",
        "J. Ho, A. Jain, and P. Abbeel, “Denoising Diffusion Probabilistic Models,” NeurIPS, 2020.",
        "ACE-Step Team, “ACE-Step: A Step Towards Music Generation Foundation Model,” arXiv:2506.00045, 2025.",
    ]
    for idx, ref in enumerate(refs, 1):
        p = paragraph(doc, f"[{idx}] {ref}", size=9.5)
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.left_indent = Cm(0.6)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
